#!/usr/bin/env python3
"""docx → md 转换：五年日志（文章随笔目录）导入 KBF 的预处理。

用法：
  .python/venv/bin/python scripts/docx2md.py <源目录> <输出目录> [--maxlength 2000]

- 递归扫描源目录所有 .docx（跳过 Word 临时锁文件 ~$*）
- 每篇日志 = 一个 md：frontmatter（name + maxlength + valid_at）+ `## 原文` 一节
- 整篇一条"时间胶囊"笔记，maxlength 声明防截断（parser 读 frontmatter）
- valid_at：文件名日期（2018-06-06）优先，否则正文落款（"写于2021-5-12"）——时间窗过滤/时间线依赖它

安全约束（2026-08-01 事故后强制）：
1. 写入目标 = os.path.join(输出目录, 输出文件名)，输出文件名只从源文件 basename 派生
2. 写入前断言真实路径展开后在输出目录内，越界直接终止（绝不覆盖源文件）
3. 输出目录必须与源目录不同（拒绝原地转换）
"""
import sys
import os
import re
import datetime

def main():
    maxlength = 2000
    excludes: list[str] = []
    args = []
    argv = sys.argv[1:]
    i = 0
    while i < len(argv):
        a = argv[i]
        if a.startswith('--maxlength'):
            if '=' in a:
                maxlength = int(a.split('=')[1])
            elif i + 1 < len(argv):
                maxlength = int(argv[i + 1])
                i += 1
            else:
                print('REFUSE: --maxlength 缺少值')
                sys.exit(2)
        elif a.startswith('--exclude'):
            if '=' in a:
                excludes.append(a.split('=')[1])
            elif i + 1 < len(argv):
                excludes.append(argv[i + 1])
                i += 1
        else:
            args.append(a)
        i += 1
    if len(args) < 2:
        print(__doc__)
        sys.exit(1)
    src_root, out_dir = args[0], args[1]

    # 安全：源目录与输出目录必须不同（防止覆盖源文件）
    if os.path.realpath(src_root) == os.path.realpath(out_dir):
        print(f'REFUSE: 源目录与输出目录相同：{src_root}')
        sys.exit(2)
    out_dir_abs = os.path.realpath(out_dir)
    os.makedirs(out_dir_abs, exist_ok=True)

    from docx import Document

    files = []
    for root, _, names in os.walk(src_root):
        for n in sorted(names):
            if n.endswith('.docx') and not n.startswith('~$'):
                if excludes and any(ex in n for ex in excludes):
                    print(f'exclude: {os.path.join(root, n)}')
                    continue
                files.append(os.path.join(root, n))
    files.sort()
    if not files:
        print(f'未找到 docx 文件：{src_root}')
        sys.exit(1)

    # 输出文件名 → 源文件（查重用；输出名只从源 basename 派生）
    name_of: dict[str, str] = {}
    for f in files:
        base = os.path.splitext(os.path.basename(f))[0]
        out = f'{base}.md'
        if out in name_of:
            out = f'{os.path.basename(os.path.dirname(f))}-{base}.md'
        name_of[out] = f

    # 防御：所有输出路径展开后必须位于输出目录内
    for out_name in name_of:
        out_path = os.path.realpath(os.path.join(out_dir_abs, out_name))
        if os.path.commonpath([out_path, out_dir_abs]) != out_dir_abs:
            print(f'REFUSE: 输出路径越界 {out_path}')
            sys.exit(2)

    n_ok = n_err = 0
    for out_name, src in name_of.items():
        try:
            doc = Document(src)
            paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            if not paras:
                # 兜底：正文在文本框/艺术字里时 doc.paragraphs 读不到，
                # 直接从 document.xml 提取所有 <w:t>（按段落分组）
                import zipfile
                xml = zipfile.ZipFile(src).read('word/document.xml').decode('utf-8', errors='replace')
                paras = []
                for p in re.split(r'</w:p>', xml):
                    t = ''.join(re.findall(r'<w:t[^>]*>([^<]*)</w:t>', p)).strip()
                    if t:
                        paras.append(t)
            if not paras:
                print(f'skip (empty): {src}')
                continue
            base = os.path.splitext(out_name)[0]
            # 日期优先级：文件名/落款 → 文件修改时间（原始写作日期，用户确认）
            valid_at = extract_date(base, paras) or date_from_mtime(src)
            with open(os.path.join(out_dir_abs, out_name), 'w', encoding='utf-8') as fh:
                # 节标题 = 文件名（详情页/图书馆直接显示日志名，如"2018-06-06"）
                # category/kind 写死：journal 是用户维护的固定分类，跳过 LLM 分类（2026-08-01 用户决策）
                fh.write(
                    f'---\nname: {base}\nmaxlength: {maxlength}\n'
                    f'category: archive/journal\nkind: experience\n'
                    f'valid_at: {valid_at}\n---\n\n## {base}\n\n'
                )
                fh.write('\n\n'.join(paras))
                fh.write('\n')
            print(f'ok: {os.path.relpath(src, src_root)} -> {out_name} ({len(paras)}段, valid_at={valid_at})')
            n_ok += 1
        except Exception as e:
            print(f'ERROR: {src}: {e}')
            n_err += 1
    print(f'\n完成：{n_ok} 篇转换，{n_err} 篇失败，输出到 {out_dir_abs}')


def extract_date(base: str, paras: list[str]) -> str:
    """日志日期提取，优先级：文件名精确日期 > 文件名英文月（2020 Dec 6）/ 中文年（2021年初）> 正文落款。

    只填可靠日期（时间窗过滤/时间线依赖它，宁缺毋滥）。
    """
    # 1. 文件名精确：2018-06-06 / 2020-9-15
    m = re.search(r'(20\d{2})[-/.年](\d{1,2})(?:[-/.月](\d{1,2}))?', base)
    if m:
        y, mo, d = m.group(1), int(m.group(2)), int(m.group(3) or 1)
        return f'{y}-{mo:02d}-{d:02d}'
    # 2. 文件名英文月：2020 Dec 6
    m = re.search(r'(20\d{2})[ _-]?([A-Za-z]{3})[. ]+(\d{1,2})', base)
    if m:
        months = {'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
                  'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12}
        mo = months.get(m.group(2).lower())
        if mo:
            return f"{m.group(1)}-{mo:02d}-{int(m.group(3)):02d}"
    # 3. 文件名中文年：2021年初 → 2021-01-01（近似）
    m = re.search(r'(20\d{2})年', base)
    if m:
        return f'{m.group(1)}-01-01'
    # 4. 正文落款：写于2021-5-12 / 2018.6.6（前后段落）
    for p in paras[:6] + paras[-3:]:
        m = re.search(r'(20\d{2})[-/.年](\d{1,2})[-/.月](\d{1,2})', p)
        if m:
            y, mo, d = m.group(1), int(m.group(2)), int(m.group(3))
            return f'{y}-{mo:02d}-{d:02d}'
    return ''


def date_from_mtime(path: str) -> str:
    """文件修改时间兜底（2026-08-01 用户决策）：无文件名/落款日期时用 docx 原始修改日期。"""
    return datetime.datetime.fromtimestamp(os.path.getmtime(path)).strftime('%Y-%m-%d')

if __name__ == '__main__':
    main()
